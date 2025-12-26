#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
특허 출원서를 마크다운에서 워드 문서(.docx)로 변환하는 스크립트
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def read_markdown_file(filename):
    """마크다운 파일 읽기"""
    with open(filename, 'r', encoding='utf-8') as f:
        return f.read()

def create_patent_document(markdown_content):
    """워드 문서 생성"""
    doc = Document()
    
    # 문서 제목 설정
    title = doc.add_heading('특허 출원 명세서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 128)
    
    # 부제
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('브라우저 기반 인공지능 자동화 단편 영상 생성 시스템 및 방법')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_paragraph()  # 빈 줄
    
    # 마크다운 파싱
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 빈 줄 건너뛰기
        if not line:
            i += 1
            continue
        
        # 수평선 건너뛰기
        if line.startswith('---'):
            i += 1
            continue
        
        # 제목 1 (## 제목)
        if line.startswith('## '):
            heading_text = line.replace('## ', '').strip('【】')
            heading = doc.add_heading(heading_text, level=1)
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.color.rgb = RGBColor(0, 70, 140)
            i += 1
            continue
        
        # 제목 2 (### 제목)
        if line.startswith('### '):
            heading_text = line.replace('### ', '').strip('【】')
            heading = doc.add_heading(heading_text, level=2)
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.color.rgb = RGBColor(0, 100, 200)
            i += 1
            continue
        
        # 제목 3 (#### 제목)
        if line.startswith('#### '):
            heading_text = line.replace('#### ', '')
            heading = doc.add_heading(heading_text, level=3)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].font.color.rgb = RGBColor(50, 120, 200)
            i += 1
            continue
        
        # 볼드 처리된 단락
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.font.bold = True
            run.font.size = Pt(11)
            i += 1
            continue
        
        # 번호 매기기 리스트
        if re.match(r'^\d+\.', line):
            p = doc.add_paragraph(line, style='List Number')
            p.runs[0].font.size = Pt(11)
            i += 1
            continue
        
        # 불릿 리스트
        if line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.runs[0].font.size = Pt(11)
            i += 1
            continue
        
        # 일반 단락
        if line and not line.startswith('#'):
            p = doc.add_paragraph()
            
            # 인라인 볼드 처리 (**텍스트**)
            text = line
            parts = re.split(r'(\*\*.*?\*\*)', text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.strip('*'))
                    run.font.bold = True
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
            
            i += 1
            continue
        
        i += 1
    
    return doc

def main():
    """메인 함수"""
    print("🔄 특허 출원서 워드 문서 생성 시작...")
    
    # 1. 특허 명세서
    print("📄 특허 명세서 변환 중...")
    md_content_app = read_markdown_file('PATENT_APPLICATION.md')
    doc_app = create_patent_document(md_content_app)
    doc_app.save('public/downloads/특허출원서_명세서.docx')
    print("✅ 특허출원서_명세서.docx 생성 완료!")
    
    # 2. 장애인 무료 출원 가이드
    print("📄 장애인 무료 출원 가이드 변환 중...")
    md_content_dis = read_markdown_file('PATENT_DISABILITY_GUIDE.md')
    doc_dis = create_patent_document(md_content_dis)
    doc_dis.save('public/downloads/장애인_무료출원_가이드.docx')
    print("✅ 장애인_무료출원_가이드.docx 생성 완료!")
    
    # 3. 특허 출원 완전 가이드
    print("📄 특허 출원 완전 가이드 변환 중...")
    md_content_guide = read_markdown_file('PATENT_GUIDE.md')
    doc_guide = create_patent_document(md_content_guide)
    doc_guide.save('public/downloads/특허출원_완전가이드.docx')
    print("✅ 특허출원_완전가이드.docx 생성 완료!")
    
    # 4. 간편 출원 가이드
    print("📄 간편 출원 가이드 변환 중...")
    doc_simple = Document()
    
    # 제목
    title = doc_simple.add_heading('특허 간편 출원 가이드', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    with open('PATENT_SIMPLE_SUBMISSION.txt', 'r', encoding='utf-8') as f:
        simple_content = f.read()
    
    # 내용 추가
    for line in simple_content.split('\n'):
        if line.strip():
            if line.startswith('='):
                continue
            elif re.match(r'^\d+\.', line):
                doc_simple.add_paragraph(line, style='List Number')
            elif line.startswith('【'):
                p = doc_simple.add_heading(line.strip('【】'), level=1)
            else:
                p = doc_simple.add_paragraph(line)
                p.runs[0].font.size = Pt(11)
    
    doc_simple.save('public/downloads/간편출원_가이드.docx')
    print("✅ 간편출원_가이드.docx 생성 완료!")
    
    print("\n🎉 모든 워드 문서 생성 완료!")
    print("\n📦 생성된 파일 목록:")
    print("  1. 특허출원서_명세서.docx")
    print("  2. 장애인_무료출원_가이드.docx")
    print("  3. 특허출원_완전가이드.docx")
    print("  4. 간편출원_가이드.docx")
    print("\n📂 저장 위치: public/downloads/")

if __name__ == '__main__':
    main()
